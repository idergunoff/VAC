import datetime
import os
import sys
from copy import deepcopy
import pickle

import pandas as pd
import numpy as np
import struct
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scipy.interpolate import interp1d
from scipy.optimize import curve_fit

from PyQt5 import uic
from PyQt5.QtWidgets import QApplication, QFileDialog, QCheckBox
from VAC_dialog import *

import pyqtgraph as pg

import matplotlib.pyplot as plt


app = QtWidgets.QApplication(sys.argv)
MainWindow = QtWidgets.QMainWindow()
ui = Ui_MainWindow()
ui.setupUi(MainWindow)
MainWindow.show()
MainWindow.showMaximized()

# pd.options.mode.chained_assignment = None  # default='warn' отключение предупреждений

all_signals = pd.DataFrame()
all_signals_old = pd.DataFrame()
columns = ['name', 'useful_depth', 'max_min', 'coeff_A', 'coeff_B', 'k_damp', 'max_h']
all_stat = pd.DataFrame(index=list(range(1, 12)), columns=columns)
all_stat_old = pd.DataFrame(index=list(range(1, 16)), columns=columns)
int_undefine_cement = []
ui.dateEdit.setDate(datetime.datetime.today() - datetime.timedelta(days=1))
name_izm = ['55-1', '53-1', '25-1', '23-1', '55-2', '53-2', '25-2', '23-2', '55-3', '53-3', '25-3', '23-3']


def open_dir():
    """ функция выбора директории с файлами измерений в формате *.lvm """
    global all_signals
    ui.progressBar.reset()
    ui.label_info.setText('Загрузка данных нового оборудования...')
    ui.label_info.setStyleSheet('color:blue')
    all_signals = pd.DataFrame()  # при открытии новой директории таблица all_signals очищается
    dir_name = QFileDialog.getExistingDirectory()  # окно выбора папки
    all_files = list()
    n = 1
    for dir_i in os.walk(dir_name):  # перебор директорий
        for file in sorted(dir_i[-1], reverse=True):  # перебор файлов в директории с обратной сортировкой
            if file.endswith('_5150m_c.lvm'):  # выбраем и заносим в список только файлы с расширением *.lvm
                all_files.append(file)
                signal = pd.read_table(os.path.join(dir_i[0], file), delimiter='\t', header=None)  # считываем таблицу
                all_signals['depth'] = signal[1]
                if n == 9:  # если записаны 8 файлов, последний файл записываем в шум и создаем колонку с глубиной
                    all_signals['noise'] = np.abs(signal[2])
                else:
                    all_signals[n] = np.abs(signal[2])
                    n_max = ((np.diff(np.sign(np.diff(all_signals[n]))) < 0).nonzero()[
                                 0] + 1).tolist()  # находим максимумы
                    n_max.insert(0, 0)  # добавляем начальную и конечную точки                              # сигнала
                    # n_max = n_max[::3]
                    n_max.append(int(len(all_signals[n])) - 1)
                    max_x = all_signals.loc[n_max, 'depth']  # выбираем x и y максимумов
                    max_y = all_signals.loc[n_max, n]
                    envelope = interp1d(max_x, max_y,
                                        kind='linear')  # выполняем интерполяцию по максимуму, получая огибающую
                    all_signals[str(n) + '_envelope'] = envelope(all_signals['depth'])
                    all_signals[str(n) + '_envelope'] = all_signals[str(n) + '_envelope'].rolling(80, min_periods=1,
                                                                                                  center=True).mean()  # усреднение
                    all_stat['name'][n] = file  # записываем названия файлов в таблицу статистики
                    all_stat['max_h'][n] = all_signals['depth'].max()
                n += 1

    ui.label_direct.setText("<b>" + dir_name + ":</b>")
    ui.label_info.setText('Загружены данные нового оборудования!')
    ui.label_info.setStyleSheet('color:green')
    checkboxes = ui.new_device.findChildren(QCheckBox)
    for item in checkboxes:
        n_izm = int(item.text()[0:2])
        if n_izm < 9:
            item.setText(str(n_izm)+' '+all_files[n_izm-1][0:20])  # присваиваем чекбоксам названия файлов
            item.setEnabled(True)

    ui.checkBox_noise.setText(all_files[8][0:20])
    number_skv = all_files[0].split('_')[0]
    ui.lineEdit_number_skv.setText(number_skv)

    ui.checkBox_noise.setEnabled(True)
    ui.pushButton_sum1.setEnabled(True)
    ui.pushButton_sum2.setEnabled(True)
    ui.pushButton_sum3.setEnabled(True)

    set_int(all_signals)

    for i in range(9, 12):      # пустые столбцы для суммы сигналов
        all_signals[str(i) + '_envelope'] = 0

    calc_to_int()  # запускаем функцию обработки сигнала в интервале


def open_dir_old():
    """ функция выбора директории с файлами измерений в формате *.TWF """
    global all_signals_old
    ui.progressBar.reset()
    ui.label_info.setText('Загрузка данных старого оборудования...')
    ui.label_info.setStyleSheet('color:blue')
    all_signals_old = pd.DataFrame()  # при открытии новой директории таблица all_signals_old очищается
    dir_name = QFileDialog.getExistingDirectory()  # окно выбора папки
    all_files = list()
    n = 1
    for file in os.listdir(dir_name):  # перебор файлов в папке
        if file.endswith('.TWF'):  # выбраем и заносим в список только файлы с расширением *.TWF
            all_files.append(file)
            f = open(os.path.join(dir_name, file), 'rb')    # открываем бинарный файл для чтения
            signal_b = f.read()[130:]                       # считываем байты кроме первых 130
            len_signal = int(len(signal_b) / 4)             # количество чисел
            uiat_b = '<' + str(len_signal) + 'i'          # формат файла для чтения
            signal = struct.unpack(uiat_b, signal_b)      # пересчет байтов в кортеж сигнала
            f.close()  # закрываем файл
            max_gl = len_signal*0.515
            if n == 1:  # первый файл записываем в шум и создаем колонку с глубиной :todo
                signal_f = interp1d(np.linspace(0, max_gl, len_signal, endpoint=True), signal, kind='linear')
                all_signals_old['depth'] = np.arange(0, max_gl, 0.05839)  # шкала глубин через 0.515 метров
                all_signals_old['noise'] = signal_f(all_signals_old['depth'])
            else:
                signal_f = interp1d(np.linspace(0, max_gl, len_signal, endpoint=True), signal, kind='linear')
                all_signals_old[n-1] = signal_f(all_signals_old['depth'])
                all_stat_old['name'][n-1] = file  # записываем названия файлов в таблицу статистики
            n += 1
            if n == 14:     # если считано 13 файлов выходим из цикла
                break
    print(all_signals_old)
    col_max = all_signals_old.columns[all_signals_old.loc[0] == all_signals_old.loc[0].max()].tolist()[0]
    k_oldtonew = all_signals_old[col_max].max()/5
    all_signals_old['noise'] = all_signals_old['noise']/k_oldtonew
    for i in range(1, len(all_signals_old.loc[0])-1):
        all_signals_old[i] = all_signals_old[i]/k_oldtonew
    print(all_signals_old)

    ui.label_direct_old.setText("<b>" + dir_name + ":</b>")
    ui.label_info.setText('Загружены данные старого оборудования!')
    ui.label_info.setStyleSheet('color:green')
    checkboxes = ui.old_device.findChildren(QCheckBox)
    for item in checkboxes:
        n_izm = int(item.text()[0:2])
        if n_izm < 13:
            item.setText(str(n_izm)+' '+all_files[n_izm][0:7])   # присваиваем чекбоксам названия файлов - первые 7 символов
            item.setEnabled(True)

    ui.checkBox_noise_old.setText(all_files[0][0:7])
    ui.checkBox_noise_old.setEnabled(True)
    ui.pushButton_sum1_old.setEnabled(True)
    ui.pushButton_sum2_old.setEnabled(True)
    ui.pushButton_sum3_old.setEnabled(True)

    set_int(all_signals_old)

    for i in range(13, 16):      # пустые столбцы для суммы сигналов
        all_signals_old[i] = 0

    calc_to_int_old()  # запускаем функцию обработки сигнала в интервале


def open_sig():
    global all_signals
    file_name = QFileDialog.getOpenFileName(filter='*.lvm')
    checkboxes = ui.new_device.findChildren(QCheckBox)
    n_chek = 0
    for item in checkboxes:
        if item.isEnabled() and int(item.text()[0:2]) not in [9, 10, 11]:
            n_chek += 1
    if n_chek == 8:
        all_signals = pd.DataFrame()
        for item in checkboxes:
            item.setEnabled(False)
            item.setStyleSheet('background: white')
        ui.pushButton_sum1.setEnabled(False)
        ui.pushButton_sum2.setEnabled(False)
        ui.pushButton_sum3.setEnabled(False)
    signal = pd.read_table(file_name[0], delimiter='\t', header=None)  # считываем таблицу
    if 'depth' in all_signals:
        if all_signals['depth'].max() < signal[1].max():
            ui.progressBar.setMaximum(len(signal[1])-len(all_signals['depth']))
            pb = 1
            ui.label_info.setText('Загружаемый сигнал больше загруженных ранее, наращиваем таблицу...')
            ui.label_info.setStyleSheet('color:blue')
            while len(all_signals['depth']) < len(signal[1]):
                all_signals.loc[len(all_signals['depth'])] = np.NaN
                ui.progressBar.setValue(pb)
                pb += 1
            all_signals['depth'] = signal[1]
            set_int(all_signals)
            ui.label_info.setText('Готово!')
            ui.label_info.setStyleSheet('color:green')
    else:
        all_signals['depth'] = signal[1]
        set_int(all_signals)
    for item in checkboxes:
        if not item.isEnabled() and int(item.text()[0:2]) not in [9, 10, 11]:
            n_sig = int(item.text()[0:2])
            item.setText(str(n_sig) + ' ' + file_name[0].split('/')[-1][:20])
            all_signals[n_sig] = np.abs(signal[2])
            n_max = ((np.diff(np.sign(np.diff(all_signals[n_sig]))) < 0).nonzero()[0] + 1).tolist()  # находим максимумы
            n_max.insert(0, 0)  # добавляем начальную и конечную точки                              # сигнала
            n_max.append(int(len(all_signals[n_sig])) - 1)
            max_x = all_signals.loc[n_max, 'depth']  # выбираем x и y максимумов
            max_y = all_signals.loc[n_max, n_sig]
            envelope = interp1d(max_x, max_y, kind='linear')  # выполняем интерполяцию по максимуму, получая огибающую
            all_signals[str(n_sig) + '_envelope'] = envelope(all_signals['depth'])
            all_signals[str(n_sig) + '_envelope'] = all_signals[str(n_sig) + '_envelope'].rolling(80, min_periods=1,
                                                                                  center=True).mean()  # усреднение
            all_stat['name'][n_sig] = item.text()
            all_stat['max_h'][n_sig] = signal[1].max()
            item.setEnabled(True)
            item.setCheckState(2)
            calc_to_int()
            break


def open_sig_old():
    global all_signals_old
    checkboxes = ui.old_device.findChildren(QCheckBox)
    for item in checkboxes:
        if not item.isEnabled() and int(item.text()[0:2]) not in [13, 14, 15]:
            print(item.text())


def calc_to_int():
    global int_min, int_max

    coeff_norm = ui.doubleSpinBox_coeff_norm.value()  # считываем значения коэффициентов из спинбоксов
    coeff_func = ui.doubleSpinBox_coeff_func.value()
    coeff_dif_res = ui.spinBox_coeff_dif_res.value()

    int_min = all_signals.index[all_signals['depth'] ==  # определяем интервалы индексов min - max
                                get_nearest_value(all_signals['depth'], ui.doubleSpinBox_int_min.value())].tolist()[0]
    int_max = all_signals.index[all_signals['depth'] ==
                                get_nearest_value(all_signals['depth'], ui.doubleSpinBox_int_max.value())].tolist()[0]
    # расчет рекомендуемого окна усреднения по формуле из Щелкуна, установка значения в спинбок
    auto_mean_win = (int_max - int_min) / (1 + 3.322 * np.log(int_max - int_min)) / 2.302585
    ui.spinBox_mean_win.setValue(auto_mean_win)
    mean_win = ui.spinBox_mean_win.value()

    ui.progressBar.setMaximum(8)
    ui.label_info.setText('Расчет кривых нового оборудования...')
    ui.label_info.setStyleSheet('color:blue')

    for i in range(1, 12):
        if (str(i) + '_envelope') in all_signals:
            if all_signals[str(i) + '_envelope'][1] != 0:
                int_max = calc_max_int(all_signals, all_stat, i)
                max_min = all_signals[str(i) + '_envelope'].iloc[int_min:int_max].max() / \
                          all_signals[str(i) + '_envelope'].iloc[int_min:int_max].min()
                all_stat['max_min'][i] = max_min  # расчет отношения макс к мин и заносим значение в таблицу статистики
                all_signals[str(i) + '_envelope-1'] = all_signals[str(i) + '_envelope'].shift(-1)  # сдвиг огибающ сигнала на 1
                all_signals[str(i) + '_mean'] = all_signals[str(i) + '_envelope'].rolling(mean_win, min_periods=1,
                                                                                          center=True).mean()  # усреднение
                all_signals[str(i) + '_norm'] = all_signals[str(i) + '_envelope'] / (all_signals[str(i) + '_mean'] + coeff_norm)
                # расчет цементограммы по усредняющей с коэффициентом
                popt, pcov = curve_fit(func1, all_signals['depth'].iloc[int_min:int_max],  # расчет коэф для функции
                                       all_signals[str(i) + '_envelope'].iloc[int_min:int_max])
                all_stat['coeff_A'][i] = popt[0]
                all_stat['coeff_B'][i] = popt[1]
                all_signals[str(i) + '_func'] = func1(all_signals['depth'], popt[0], popt[1]) + coeff_func  # построение
                # фунции в интервале сигнала со сдвигом по Y на coeff_func
                """Расчет коэффициента затухания"""
                all_signals[str(i) + '_func_kdamp'] = all_signals[str(i) + '_func'][0]/all_signals[str(i) + '_func']
                i_kdamp = all_signals.index[all_signals[str(i) + '_func_kdamp'] == get_nearest_value(all_signals[str(i) +
                                                                                     '_func_kdamp'], np.exp(1))].tolist()[0]
                # all_stat['k_damp'][i] = 0.2029*np.log(np.exp(1)/all_signals['depth'][i_kdamp])+0.9127
                all_stat['k_damp'][i] = np.exp(1) / all_signals['depth'][i_kdamp]
                # расчет
                # коэффициента затухания и пресчет в размерность коэффициента качества
                all_signals[str(i) + '_rel_ampl'] = (all_signals[str(i) + '_envelope-1'] - all_signals[str(i) + '_func']) /\
                                                    all_signals[str(i) + '_envelope-1']  # расчет цементограммы по функции
                all_signals[str(i) + '_mean-1'] = all_signals[str(i) + '_mean'].shift(-1)  # сдвиг усредненного сигнала на 1
                all_signals[str(i) + '_diff_norm'] = (all_signals[str(i) + '_mean-1'] - all_signals[str(i) + '_mean']) / \
                                                     all_signals[str(i) + '_mean']  # нормированная производная
                all_signals[str(i) + '_diff_norm-1'] = all_signals[str(i) + '_diff_norm'].shift(-1)  # сдвиг
                all_signals[str(i) + '_diff_result'] = (all_signals[str(i) + '_diff_norm-1'] + all_signals[
                    str(i) + '_diff_norm']) * coeff_dif_res  # расчет цементограммы по производной (как в щелкуне)

                std_25 = all_signals[str(i) + '_envelope'].iloc[-430:].std()  #стандартное отклонение по последним 25 метрам
                for n, k in enumerate(all_signals[str(i) + '_mean'].iloc[int_min:int_max]):
                    if k <= std_25 + all_signals[str(i) + '_mean'].iloc[int_min:int_max].min():
                        useful_depth = all_signals['depth'][n]
                        all_stat['useful_depth'][i] = useful_depth
                        break
                ui.progressBar.setValue(i)

    checkboxes = ui.new_device.findChildren(QCheckBox)
    for item in checkboxes:
        if item.isEnabled():
            n_izm = int(item.text()[0:2])
            item.setToolTip('Коэффициент затухания - '+str(round(all_stat['k_damp'][n_izm], 3))+
                             '\nЭффективная глубина - ' + str(round(all_stat['useful_depth'][n_izm], 2))+
                             ' м\nМакс/Мин - '+str(round(all_stat['max_min'][n_izm], 3)))
            item.setStyleSheet(k_damp_detect('new', n_izm))

    all_signals.to_csv('result.txt', sep='\t')
    all_stat.to_csv('all_stat.txt', sep='\t')
    ui.label_info.setText('Готово!')
    ui.label_info.setStyleSheet('color:green')
    choice_signal()


def calc_to_int_old():
    global int_min_old, int_max_old

    coeff_norm = ui.doubleSpinBox_coeff_norm_old.value()  # считываем значения коэффициентов из спинбоксов
    coeff_func = ui.doubleSpinBox_coeff_func_old.value()
    coeff_dif_res = ui.spinBox_coeff_dif_res_old.value()

    int_min_old = all_signals_old.index[all_signals_old['depth'] ==  # определяем интервалы индексов min - max
                            get_nearest_value(all_signals_old['depth'], ui.doubleSpinBox_int_min.value())].tolist()[0]
    int_max_old = all_signals_old.index[all_signals_old['depth'] ==
                            get_nearest_value(all_signals_old['depth'], ui.doubleSpinBox_int_max.value())].tolist()[0]
    auto_mean_win = (int_max_old - int_min_old) / (1 + 3.322 * np.log(int_max_old - int_min_old)) / 2.302585
    ui.spinBox_mean_win_old.setValue(auto_mean_win)
    mean_win = ui.spinBox_mean_win_old.value()
    ui.progressBar.setMaximum(12)
    ui.label_info.setText('Расчет кривых старого оборудования...')
    ui.label_info.setStyleSheet('color:blue')

    for i in range(1, 16):
        if all_signals_old[i][1] != 0:
            max_min = all_signals_old[i].iloc[int_min_old:int_max_old].max() / \
                      all_signals_old[i].iloc[int_min_old:int_max_old].min()
            all_stat_old['max_min'][i] = max_min  # расчет отношения макс к мин и заносим значение в таблицу статистики
            all_signals_old[str(i) + '-1'] = all_signals_old[i].shift(-1)  # сдвиг огибающ сигнала на 1
            all_signals_old[str(i) + '_mean'] = all_signals_old[i].rolling(mean_win, min_periods=1, center=True).mean()
            all_signals_old[str(i) + '_norm'] = all_signals_old[i] / (all_signals_old[str(i) + '_mean'] + coeff_norm)
            # расчет цементограммы по усредняющей с коэффициентом
            popt, pcov = curve_fit(func1, all_signals_old['depth'].iloc[int_min_old:int_max_old],
                               all_signals_old[i].iloc[int_min_old:int_max_old])    # расчет коэф для функции
            all_stat_old['coeff_A'][i] = popt[0]
            all_stat_old['coeff_B'][i] = popt[1]
            all_signals_old[str(i) + '_func'] = func1(all_signals_old['depth'], popt[0], popt[1]) + coeff_func
            # построение фунции в интервале сигнала со сдвигом по Y на coeff_func
            """Расчет коэффициента затухания"""
            all_signals_old[str(i) + '_func_kdamp'] = all_signals_old[str(i) + '_func'][0] / \
                                                      all_signals_old[str(i) + '_func']
            i_kdamp = all_signals_old.index[
                all_signals_old[str(i) + '_func_kdamp'] == get_nearest_value(all_signals_old[str(i) + '_func_kdamp'],
                                                                         np.exp(1))].tolist()[0]
            all_stat_old['k_damp'][i] = 0.2029 * np.log(np.exp(1) / all_signals_old['depth'][i_kdamp]) + 0.9127
            all_stat_old['k_damp'][i] = np.exp(1) / all_signals_old['depth'][i_kdamp]
            all_signals_old[str(i) + '_rel_ampl'] = (all_signals_old[str(i) + '-1'] -  # расчет цементограммы по функции
                                                     all_signals_old[str(i) + '_func']) / all_signals_old[str(i) + '-1']
            all_signals_old[str(i) + '_mean-1'] = all_signals_old[str(i) + '_mean'].shift(-1)
                                                                                        # сдвиг усредненного сигнала
            all_signals_old[str(i) + '_diff_norm'] = (all_signals_old[str(i) + '_mean-1'] -  # нормированная производная
                                                  all_signals_old[str(i) + '_mean']) / all_signals_old[str(i) + '_mean']
            all_signals_old[str(i) + '_diff_norm-1'] = all_signals_old[str(i) + '_diff_norm'].shift(-1)  # сдвиг
            all_signals_old[str(i) + '_diff_result'] = (all_signals_old[str(i) + '_diff_norm-1'] + all_signals_old[
                str(i) + '_diff_norm']) * coeff_dif_res  # расчет цементограммы по производной (как в щелкуне)

            std_25 = all_signals_old[i].iloc[-50:].std()  # стандартное отклонение по последним 25 метрам
            for n, k in enumerate(all_signals_old[str(i) + '_mean'].iloc[int_min_old:int_max_old]):
                if k <= std_25 + all_signals_old[str(i) + '_mean'].iloc[int_min_old:int_max_old].min():
                    useful_depth = all_signals_old['depth'][n]
                    all_stat_old['useful_depth'][i] = useful_depth
                    break
            ui.progressBar.setValue(i)
    # расчет рекомендуемого окна усреднения по формуле из Щелкуна, установка значения в спинбок

    print(all_stat_old)
    print(all_signals_old)
    checkboxes = ui.old_device.findChildren(QCheckBox)
    for item in checkboxes:
        n_izm = int(item.text()[0:2])
        item.setToolTip('Коэффициент затухания - ' + str(round(all_stat_old['k_damp'][n_izm], 3)) +
                         '\nЭффективная глубина - ' + str(round(all_stat_old['useful_depth'][n_izm], 2)) +
                         ' м\nМакс/Мин - ' + str(round(all_stat_old['max_min'][n_izm], 3)))
        item.setStyleSheet(k_damp_detect('old', n_izm))

    ui.label_info.setText('Готово!')
    ui.label_info.setStyleSheet('color:green')
    all_signals_old.to_csv('result_old.txt', sep='\t')
    all_stat_old.to_csv('all_stat_old.txt', sep='\t')
    choice_signal()


def calc():
    if ui.checkBox_signal1.isEnabled():
        calc_to_int()
    if ui.checkBox_signal_old_1.isEnabled():
        calc_to_int_old()


def choice_signal():
    """
    выбор сигналов для построения
    """
    ui.graphicsView.clear()
    if ui.checkBox_signal1.checkState() == 2 and ui.checkBox_signal1.isEnabled():
        plot_graph(1, 1.8, [1], 0)
    if ui.checkBox_signal2.checkState() == 2 and ui.checkBox_signal2.isEnabled():
        plot_graph(2, 1.8, [3, 2], 0)
    if ui.checkBox_signal3.checkState() == 2 and ui.checkBox_signal3.isEnabled():
        plot_graph(3, 1.5, [6, 4], 0)
    if ui.checkBox_signal4.checkState() == 2 and ui.checkBox_signal4.isEnabled():
        plot_graph(4, 1.5, [10, 5, 1, 5], 0)
    if ui.checkBox_signal5.checkState() == 2 and ui.checkBox_signal5.isEnabled():
        plot_graph(5, 1.8, [1], 50)
    if ui.checkBox_signal6.checkState() == 2 and ui.checkBox_signal6.isEnabled():
        plot_graph(6, 1.8, [3, 2], 50)
    if ui.checkBox_signal7.checkState() == 2 and ui.checkBox_signal7.isEnabled():
        plot_graph(7, 1.5, [6, 4], 50)
    if ui.checkBox_signal8.checkState() == 2 and ui.checkBox_signal8.isEnabled():
        plot_graph(8, 1.5, [10, 5, 1, 5], 50)
    if ui.checkBox_sum1.checkState() == 2 and ui.checkBox_sum1.isEnabled():
        plot_graph(9, 2.5, [3, 3], 75)
    if ui.checkBox_sum2.checkState() == 2 and ui.checkBox_sum2.isEnabled():
        plot_graph(10, 2.5, [3, 3], 75)
    if ui.checkBox_sum3.checkState() == 2 and ui.checkBox_sum3.isEnabled():
        plot_graph(11, 2.5, [3, 3], 75)

    if ui.checkBox_signal_old_1.checkState() == 2 and ui.checkBox_signal_old_1.isEnabled():
        plot_graph_old(1, 1.8, [1], 100)
    if ui.checkBox_signal_old_2.checkState() == 2 and ui.checkBox_signal_old_2.isEnabled():
        plot_graph_old(2, 1.8, [3, 2], 100)
    if ui.checkBox_signal_old_3.checkState() == 2 and ui.checkBox_signal_old_3.isEnabled():
        plot_graph_old(3, 1.5, [6, 4], 100)
    if ui.checkBox_signal_old_4.checkState() == 2 and ui.checkBox_signal_old_4.isEnabled():
        plot_graph_old(4, 1.5, [10, 5, 1, 5], 100)
    if ui.checkBox_signal_old_5.checkState() == 2 and ui.checkBox_signal_old_5.isEnabled():
        plot_graph_old(5, 1.8, [1], 150)
    if ui.checkBox_signal_old_6.checkState() == 2 and ui.checkBox_signal_old_6.isEnabled():
        plot_graph_old(6, 1.8, [3, 2], 150)
    if ui.checkBox_signal_old_7.checkState() == 2 and ui.checkBox_signal_old_7.isEnabled():
        plot_graph_old(7, 1.5, [6, 4], 150)
    if ui.checkBox_signal_old_8.checkState() == 2 and ui.checkBox_signal_old_8.isEnabled():
        plot_graph_old(8, 1.5, [10, 5, 1, 5], 150)
    if ui.checkBox_signal_old_9.checkState() == 2 and ui.checkBox_signal_old_9.isEnabled():
        plot_graph_old(9, 1.8, [1], 200)
    if ui.checkBox_signal_old_10.checkState() == 2 and ui.checkBox_signal_old_10.isEnabled():
        plot_graph_old(10, 1.8, [3, 2], 200)
    if ui.checkBox_signal_old_11.checkState() == 2 and ui.checkBox_signal_old_11.isEnabled():
        plot_graph_old(11, 1.5, [6, 4], 200)
    if ui.checkBox_signal_old_12.checkState() == 2 and ui.checkBox_signal_old_12.isEnabled():
        plot_graph_old(12, 1.5, [10, 5, 1, 5], 200)
    if ui.checkBox_sum1_old.checkState() == 2 and ui.checkBox_sum1_old.isEnabled():
        plot_graph_old(13, 2.5, [3, 3], 175)
    if ui.checkBox_sum2_old.checkState() == 2 and ui.checkBox_sum2_old.isEnabled():
        plot_graph_old(14, 2.5, [3, 3], 175)
    if ui.checkBox_sum3_old.checkState() == 2 and ui.checkBox_sum3_old.isEnabled():
        plot_graph_old(15, 2.5, [3, 3], 175)

    if ui.checkBox_noise.checkState() == 2 and ui.checkBox_noise.isEnabled():
        ui.graphicsView.plot(x=all_signals['depth'].iloc[int_min:int_max].tolist(),
                               y=all_signals['noise'].iloc[int_min:int_max].tolist(),
                               pen=pg.mkPen(color='c', dash=[2, 1, 4], width=1))
    if ui.checkBox_noise_old.checkState() == 2 and ui.checkBox_noise_old.isEnabled():
        ui.graphicsView.plot(x=all_signals_old['depth'].iloc[int_min_old:int_max_old].tolist(),
                               y=all_signals_old['noise'].iloc[int_min_old:int_max_old].tolist(),
                               pen=pg.mkPen(color='c', dash=[2, 1, 4], width=1))


def plot_graph(n_sig, sig_width, sig_dash, color):
    """
    выбор и построение кривых для выбранных сигналов
    принимает параметры линии для построения
    """
    ui.graphicsView.showGrid(x=True, y=True)  # грид-сетка
    if ui.checkBox_origin_sig.checkState() == 2:
        int_max = calc_max_int(all_signals, all_stat, n_sig)
        ui.graphicsView.plot(x=all_signals['depth'].iloc[int_min:int_max].tolist(),
                               y=all_signals[n_sig].iloc[int_min:int_max].tolist(),
                               pen=pg.mkPen(width=sig_width, dash=sig_dash))
    if ui.checkBox_envelop.checkState() == 2:
        int_max = calc_max_int(all_signals, all_stat, n_sig)
        ui.graphicsView.plot(x=all_signals['depth'].iloc[int_min:int_max].tolist(),
                               y=all_signals[str(n_sig) + '_envelope'].iloc[int_min:int_max].tolist(),
                               pen=pg.mkPen(color=(255, 0, color), width=sig_width, dash=sig_dash))
    if ui.checkBox_func.checkState() == 2:
        int_max = calc_max_int(all_signals, all_stat, n_sig)
        ui.graphicsView.plot(x=all_signals['depth'].iloc[int_min:int_max].tolist(),
                               y=all_signals[str(n_sig) + '_func'].iloc[int_min:int_max].tolist(),
                               pen=pg.mkPen(color=(0, 255, color), width=sig_width, dash=sig_dash))
    if ui.checkBox_mean.checkState() == 2:
        int_max = calc_max_int(all_signals, all_stat, n_sig)
        ui.graphicsView.plot(x=all_signals['depth'].iloc[int_min:int_max].tolist(),
                               y=all_signals[str(n_sig) + '_mean'].iloc[int_min:int_max].tolist(),
                               pen=pg.mkPen(color=(100, 100, 255 - color), width=sig_width, dash=sig_dash))
    if ui.checkBox_norm.checkState() == 2:
        int_max = calc_max_int(all_signals, all_stat, n_sig)
        ui.graphicsView.plot(x=all_signals['depth'].iloc[int_min:int_max].tolist(),
                               y=all_signals[str(n_sig) + '_norm'].iloc[int_min:int_max].tolist(),
                               pen=pg.mkPen(color=(255, 255, color), width=sig_width, dash=sig_dash))
    if ui.checkBox_rel_ampl.checkState() == 2:
        int_max = calc_max_int(all_signals, all_stat, n_sig)
        ui.graphicsView.plot(x=all_signals['depth'].iloc[int_min:int_max].tolist(),
                               y=all_signals[str(n_sig) + '_rel_ampl'].iloc[int_min:int_max].tolist(),
                               pen=pg.mkPen(color='c', width=sig_width, dash=sig_dash))
    if ui.checkBox_diff_result.checkState() == 2:
        int_max = calc_max_int(all_signals, all_stat, n_sig)
        ui.graphicsView.plot(x=all_signals['depth'].iloc[int_min:int_max].tolist(),
                               y=all_signals[str(n_sig) + '_diff_result'].iloc[int_min:int_max].tolist(),
                               pen=pg.mkPen(color=(255, 255, color), width=sig_width, dash=sig_dash))


def plot_graph_old(n_sig, sig_width, sig_dash, color):
    """
    выбор и построение кривых для выбранных сигналов
    принимает параметры линии для построения
    """
    ui.graphicsView.showGrid(x=True, y=True)  # грид-сетка
    if ui.checkBox_envelop_old.checkState() == 2:
        ui.graphicsView.plot(x=all_signals_old['depth'].iloc[int_min_old:int_max_old].tolist(),
                               y=all_signals_old[n_sig].iloc[int_min_old:int_max_old].tolist(),
                               pen=pg.mkPen(color=(255, 0, color), width=sig_width, dash=sig_dash))
    if ui.checkBox_func_old.checkState() == 2:
        ui.graphicsView.plot(x=all_signals_old['depth'].iloc[int_min_old:int_max_old].tolist(),
                               y=all_signals_old[str(n_sig) + '_func'].iloc[int_min_old:int_max_old].tolist(),
                               pen=pg.mkPen(color=(0, 255, color), width=sig_width, dash=sig_dash))
    if ui.checkBox_mean_old.checkState() == 2:
        ui.graphicsView.plot(x=all_signals_old['depth'].iloc[int_min_old:int_max_old].tolist(),
                               y=all_signals_old[str(n_sig) + '_mean'].iloc[int_min_old:int_max_old].tolist(),
                               pen=pg.mkPen(color=(100, 100, 255 - color), width=sig_width, dash=sig_dash))
    if ui.checkBox_norm_old.checkState() == 2:
        ui.graphicsView.plot(x=all_signals_old['depth'].iloc[int_min_old:int_max_old].tolist(),
                               y=all_signals_old[str(n_sig) + '_norm'].iloc[int_min_old:int_max_old].tolist(),
                               pen=pg.mkPen(color=(255, 255, color), width=sig_width, dash=sig_dash))
    if ui.checkBox_rel_ampl_old.checkState() == 2:
        ui.graphicsView.plot(x=all_signals_old['depth'].iloc[int_min_old:int_max_old].tolist(),
                               y=all_signals_old[str(n_sig) + '_rel_ampl'].iloc[int_min_old:int_max_old].tolist(),
                               pen=pg.mkPen(color='c', width=sig_width, dash=sig_dash))
    if ui.checkBox_diff_result_old.checkState() == 2:
        ui.graphicsView.plot(x=all_signals_old['depth'].iloc[int_min_old:int_max_old].tolist(),
                               y=all_signals_old[str(n_sig) + '_diff_result'].iloc[int_min_old:int_max_old].tolist(),
                               pen=pg.mkPen(color=(255, 255, color), width=sig_width, dash=sig_dash))


def change_mean_win():
    """
    пересчёт кривых при изменении коэффициентов
    """
    mean_win = ui.spinBox_mean_win.value()
    coeff_norm = ui.doubleSpinBox_coeff_norm.value()
    coeff_dif_res = ui.spinBox_coeff_dif_res.value()
    for i in range(1, 12):
        if (str(i) + '_envelope') in all_signals:
            if all_signals[str(i) + '_envelope'][1] != 0:
                all_signals[str(i) + '_mean'] = all_signals[str(i) + '_envelope'].rolling(mean_win, min_periods=1,
                                                                                          center=True).mean()
                all_signals[str(i) + '_norm'] = all_signals[str(i) + '_envelope'] / (all_signals[str(i) + '_mean'] + coeff_norm)
                all_signals[str(i) + '_mean-1'] = all_signals[str(i) + '_mean'].shift(-1)
                all_signals[str(i) + '_diff_norm'] = (all_signals[str(i) + '_mean-1'] - all_signals[str(i) + '_mean']) / \
                                                     all_signals[str(i) + '_mean']
                all_signals[str(i) + '_diff_norm-1'] = all_signals[str(i) + '_diff_norm'].shift(-1)
                all_signals[str(i) + '_diff_result'] = (all_signals[str(i) + '_diff_norm-1'] + all_signals[
                    str(i) + '_diff_norm']) * coeff_dif_res

    all_signals.to_csv('result.txt', sep='\t')
    choice_signal()


def change_mean_win_old():
    """
    пересчёт кривых при изменении коэффициентов
    """
    coeff_norm = ui.doubleSpinBox_coeff_norm_old.value()  # считываем значения коэффициентов из спинбоксов
    coeff_dif_res = ui.spinBox_coeff_dif_res_old.value()
    mean_win = ui.spinBox_mean_win_old.value()

    for i in range(1, 16):
        if all_signals_old[i][1] != 0:
            all_signals_old[str(i) + '_mean'] = all_signals_old[i].rolling(mean_win, min_periods=1, center=True).mean()
            all_signals_old[str(i) + '_norm'] = all_signals_old[i] / (all_signals_old[str(i) + '_mean'] + coeff_norm)


            all_signals_old[str(i) + '_mean-1'] = all_signals_old[str(i) + '_mean'].shift(-1)  # сдвиг усредненного сигнала
            all_signals_old[str(i) + '_diff_norm'] = (all_signals_old[str(i) + '_mean-1'] -  # нормированная производная
                                                      all_signals_old[str(i) + '_mean']) / all_signals_old[str(i) + '_mean']
            all_signals_old[str(i) + '_diff_norm-1'] = all_signals_old[str(i) + '_diff_norm'].shift(-1)  # сдвиг
            all_signals_old[str(i) + '_diff_result'] = (all_signals_old[str(i) + '_diff_norm-1'] + all_signals_old[
                str(i) + '_diff_norm']) * coeff_dif_res  # расчет цементограммы по производной (как в щелкуне)

    all_signals_old.to_csv('result_old.txt', sep='\t')
    choice_signal()


def func1(x, a, b):
    """ функция затухания """
    return a * b ** x
    # return a*(x-b)**c


def change_func():
    """
    сдвиг функции затухания
    """
    coeff_func = ui.doubleSpinBox_coeff_func.value()
    kA = ui.doubleSpinBox_kA.value()
    kB = ui.doubleSpinBox_kB.value()
    for i in range(1, 12):
        if (str(i) + '_envelope') in all_signals:
            if all_signals[str(i) + '_envelope'][1] != 0:
                all_signals[str(i) + '_func'] = func1(all_signals['depth'], all_stat['coeff_A'][i]+kA, all_stat['coeff_B'][i]+kB) + coeff_func
                all_signals[str(i) + '_rel_ampl'] = (all_signals[str(i) + '_envelope-1'] - all_signals[str(i) + '_func']) / \
                                                    all_signals[str(i) + '_envelope-1']
    choice_signal()


def change_func_old():
    """
    сдвиг функции затухания
    """
    coeff_func = ui.doubleSpinBox_coeff_func_old.value()
    kA = ui.doubleSpinBox_kA_old.value()
    kB = ui.doubleSpinBox_kB_old.value()
    for i in range(1, 16):
        if all_signals_old[i][1] != 0:
            all_signals_old[str(i) + '_func'] = func1(all_signals_old['depth'], all_stat_old['coeff_A'][i]+kA, all_stat_old['coeff_B'][i]+kB) + coeff_func
            all_signals_old[str(i) + '_rel_ampl'] = (all_signals_old[str(i) + '-1'] - all_signals_old[str(i) + '_func']) / \
                                                    all_signals_old[str(i) + '-1']  # расчет цементограммы по функции
    choice_signal()


def check_int_sig():
    """ проверка интервала сигнала """
    ui.doubleSpinBox_int_min.setMaximum(ui.doubleSpinBox_int_max.value())
    ui.doubleSpinBox_int_max.setMinimum(ui.doubleSpinBox_int_min.value())


def get_nearest_value(iterable, value):
    """ функция поиска ближайщего значения """
    return min(iterable, key=lambda x: abs(x - value))


def sum_signals(n_sum):
    n_signals = 0
    all_signals[str(n_sum) + '_envelope'] = 0
    if ui.checkBox_signal1.checkState() == 2:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope'] + all_signals['1_envelope']
        n_signals += 1
    if ui.checkBox_signal2.checkState() == 2:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope'] + all_signals['2_envelope']
        n_signals += 1
    if ui.checkBox_signal3.checkState() == 2:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope'] + all_signals['3_envelope']
        n_signals += 1
    if ui.checkBox_signal4.checkState() == 2:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope'] + all_signals['4_envelope']
        n_signals += 1
    if ui.checkBox_signal5.checkState() == 2:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope'] + all_signals['5_envelope']
        n_signals += 1
    if ui.checkBox_signal6.checkState() == 2:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope'] + all_signals['6_envelope']
        n_signals += 1
    if ui.checkBox_signal7.checkState() == 2:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope'] + all_signals['7_envelope']
        n_signals += 1
    if ui.checkBox_signal8.checkState() == 2:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope'] + all_signals['8_envelope']
        n_signals += 1
    if n_signals != 0:
        all_signals[str(n_sum) + '_envelope'] = all_signals[str(n_sum) + '_envelope']/n_signals
        coeff_norm = ui.doubleSpinBox_coeff_norm.value()  # считываем значения коэффициентов из спинбоксов
        coeff_func = ui.doubleSpinBox_coeff_func.value()
        coeff_dif_res = ui.spinBox_coeff_dif_res.value()
        mean_win = ui.spinBox_mean_win.value()
        max_min = all_signals[str(n_sum) + '_envelope'].iloc[int_min:int_max].max() / \
                  all_signals[str(n_sum) + '_envelope'].iloc[int_min:int_max].min()
        all_stat['max_min'][n_sum] = max_min  # расчет отношения макс к мин и заносим значение в таблицу статистики
        all_signals[str(n_sum) + '_envelope-1'] = all_signals[str(n_sum) + '_envelope'].shift(-1)  # сдвиг огибающ сигнала на 1
        all_signals[str(n_sum) + '_mean'] = all_signals[str(n_sum) + '_envelope'].rolling(mean_win, min_periods=1,
                                                                                  center=True).mean()  # усреднение
        all_signals[str(n_sum) + '_norm'] = all_signals[str(n_sum) + '_envelope'] / (all_signals[str(n_sum) + '_mean'] + coeff_norm)
        # расчет цементограммы по усредняющей с коэффициентом
        popt, pcov = curve_fit(func1, all_signals['depth'].iloc[int_min:int_max],  # расчет коэф для функции
                               all_signals[str(n_sum) + '_envelope'].iloc[int_min:int_max])
        all_stat['coeff_A'][n_sum] = popt[0]
        all_stat['coeff_B'][n_sum] = popt[1]
        all_signals[str(n_sum) + '_func'] = func1(all_signals['depth'], popt[0], popt[1]) + coeff_func  # построение
        # фунции в интервале сигнала со сдвигом по Y на coeff_func

        """Расчет коэффициента затухания"""
        all_signals[str(n_sum) + '_func_kdamp'] = all_signals[str(n_sum) + '_func'][0]/all_signals[str(n_sum) + '_func']
        i_kdamp = all_signals.index[all_signals[str(n_sum) + '_func_kdamp'] == get_nearest_value(all_signals[str(n_sum)
                                                                             +'_func_kdamp'], np.exp(1))].tolist()[0]
        # all_stat['k_damp'][n_sum] = 0.2029*np.log(np.exp(1)/all_signals['depth'][i_kdamp])+0.9127
        all_stat['k_damp'][n_sum] = np.exp(1) / all_signals['depth'][i_kdamp]

        all_signals[str(n_sum) + '_rel_ampl'] = (all_signals[str(n_sum) + '_envelope-1'] - all_signals[str(n_sum) + '_func']) / \
                                            all_signals[str(n_sum) + '_envelope-1']  # расчет цементограммы по функции
        all_signals[str(n_sum) + '_mean-1'] = all_signals[str(n_sum) + '_mean'].shift(-1)  # сдвиг усредненного сигнала на 1
        all_signals[str(n_sum) + '_diff_norm'] = (all_signals[str(n_sum) + '_mean-1'] - all_signals[str(n_sum) + '_mean']) / \
                                             all_signals[str(n_sum) + '_mean']  # нормированная производная
        all_signals[str(n_sum) + '_diff_norm-1'] = all_signals[str(n_sum) + '_diff_norm'].shift(-1)  # сдвиг
        all_signals[str(n_sum) + '_diff_result'] = (all_signals[str(n_sum) + '_diff_norm-1'] + all_signals[
            str(n_sum) + '_diff_norm']) * coeff_dif_res  # расчет цементограммы по производной (как в щелкуне)
        std_25 = all_signals[str(n_sum) + '_envelope'].iloc[-430:].std()  # стандартное отклонение по последним 25 метрам
        for n, k in enumerate(all_signals[str(n_sum) + '_mean'].iloc[int_min:int_max]):
            if k <= std_25 + all_signals[str(n_sum) + '_mean'].iloc[int_min:int_max].min():
                useful_depth = all_signals['depth'][n]
                all_stat['useful_depth'][n_sum] = useful_depth
                break
        print(all_stat)
        all_signals.to_csv('result.txt', sep='\t')
        all_stat.to_csv('all_stat.txt', sep='\t')
        choice_signal()


def sum_signals_old(n_sum):
    n_signals = 0
    all_signals_old[n_sum] = 0
    if ui.checkBox_signal_old_1.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[1]
        n_signals += 1
    if ui.checkBox_signal_old_2.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[2]
        n_signals += 1
    if ui.checkBox_signal_old_3.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[3]
        n_signals += 1
    if ui.checkBox_signal_old_4.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[4]
        n_signals += 1
    if ui.checkBox_signal_old_5.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[5]
        n_signals += 1
    if ui.checkBox_signal_old_6.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[6]
        n_signals += 1
    if ui.checkBox_signal_old_7.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[7]
        n_signals += 1
    if ui.checkBox_signal_old_8.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[8]
        n_signals += 1
    if ui.checkBox_signal_old_9.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[8]
        n_signals += 1
    if ui.checkBox_signal_old_10.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[10]
        n_signals += 1
    if ui.checkBox_signal_old_11.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[11]
        n_signals += 1
    if ui.checkBox_signal_old_12.checkState() == 2:
        all_signals_old[n_sum] = all_signals_old[n_sum] + all_signals_old[12]
        n_signals += 1
    if n_signals != 0:
        all_signals_old[n_sum] = all_signals_old[n_sum]/n_signals
        coeff_norm = ui.doubleSpinBox_coeff_norm_old.value()  # считываем значения коэффициентов из спинбоксов
        coeff_func = ui.doubleSpinBox_coeff_func_old.value()
        coeff_dif_res = ui.spinBox_coeff_dif_res_old.value()
        mean_win = ui.spinBox_mean_win_old.value()

        max_min = all_signals_old[n_sum].iloc[int_min_old:int_max_old].max() / \
                  all_signals_old[n_sum].iloc[int_min_old:int_max_old].min()
        all_stat_old['max_min'][n_sum] = max_min  # расчет отношения макс к мин и заносим значение в таблицу статистики
        all_signals_old[str(n_sum) + '-1'] = all_signals_old[n_sum].shift(-1)  # сдвиг огибающ сигнала на 1
        all_signals_old[str(n_sum) + '_mean'] = all_signals_old[n_sum].rolling(mean_win, min_periods=1, center=True).mean()
        all_signals_old[str(n_sum) + '_norm'] = all_signals_old[n_sum] / (all_signals_old[str(n_sum) + '_mean'] + coeff_norm)
        # расчет цементограммы по усредняющей с коэффициентом
        popt, pcov = curve_fit(func1, all_signals_old['depth'].iloc[int_min_old:int_max_old],  # расчет коэф для функции
                               all_signals_old[n_sum].iloc[int_min_old:int_max_old])
        all_stat_old['coeff_A'][n_sum] = popt[0]
        all_stat_old['coeff_B'][n_sum] = popt[1]
        all_signals_old[str(n_sum) + '_func'] = func1(all_signals_old['depth'], popt[0], popt[1]) + coeff_func  # построение
        # фунции в интервале сигнала со сдвигом по Y на coeff_func

        """Расчет коэффициента затухания"""
        all_signals_old[str(n_sum) + '_func_kdamp'] = all_signals_old[str(n_sum) + '_func'][0] / all_signals_old[
            str(n_sum) + '_func']
        i_kdamp = all_signals_old.index[
            all_signals_old[str(n_sum) + '_func_kdamp'] == get_nearest_value(all_signals_old[str(n_sum) + '_func_kdamp'],
                                                                         np.exp(1))].tolist()[0]
        # all_stat_old['k_damp'][n_sum] = 0.2029 * np.log(np.exp(1) / all_signals_old['depth'][i_kdamp]) + 0.9127
        all_stat_old['k_damp'][n_sum] = np.exp(1) / all_signals_old['depth'][i_kdamp]

        all_signals_old[str(n_sum) + '_rel_ampl'] = (all_signals_old[str(n_sum) + '-1'] - all_signals_old[str(n_sum) + '_func']) / \
                                                all_signals_old[str(n_sum) + '-1']  # расчет цементограммы по функции
        all_signals_old[str(n_sum) + '_mean-1'] = all_signals_old[str(n_sum) + '_mean'].shift(-1)  # сдвиг усредненного сигнала
        all_signals_old[str(n_sum) + '_diff_norm'] = (all_signals_old[str(n_sum) + '_mean-1'] -  # нормированная производная
                                                  all_signals_old[str(n_sum) + '_mean']) / all_signals_old[str(n_sum) + '_mean']
        all_signals_old[str(n_sum) + '_diff_norm-1'] = all_signals_old[str(n_sum) + '_diff_norm'].shift(-1)  # сдвиг
        all_signals_old[str(n_sum) + '_diff_result'] = (all_signals_old[str(n_sum) + '_diff_norm-1'] + all_signals_old[
            str(n_sum) + '_diff_norm']) * coeff_dif_res  # расчет цементограммы по производной (как в щелкуне)

        std_25 = all_signals_old[n_sum].iloc[-50:].std()  # стандартное отклонение по последним 25 метрам
        for n, k in enumerate(all_signals_old[str(n_sum) + '_mean'].iloc[int_min_old:int_max_old]):
            if k <= std_25 + all_signals_old[str(n_sum) + '_mean'].iloc[int_min_old:int_max_old].min():
                useful_depth = all_signals_old['depth'][n]
                all_stat_old['useful_depth'][n_sum] = useful_depth
                break
        print(all_stat)
        all_signals_old.to_csv('result_old.txt', sep='\t')
        all_stat_old.to_csv('all_stat_old.txt', sep='\t')
        choice_signal()


def sum1():
    all_stat['name'][9] = text_sum()
    sum_signals(9)
    ui.checkBox_sum1.setText('9 '+text_sum())
    ui.checkBox_sum1.setToolTip('Коэффициент затухания - ' + str(round(all_stat['k_damp'][9], 3)) +
                                     '\nЭффективная глубина - ' + str(round(all_stat['useful_depth'][9], 2)) +
                                     ' м\nМакс/Мин - ' + str(round(all_stat['max_min'][9], 3)))
    ui.checkBox_sum1.setStyleSheet(k_damp_detect('new', 9))
    ui.checkBox_sum1.setEnabled(True)


def sum2():
    all_stat['name'][10] = text_sum()
    sum_signals(10)
    ui.checkBox_sum2.setText('10 '+text_sum())
    ui.checkBox_sum2.setToolTip('Коэффициент затухания - ' + str(round(all_stat['k_damp'][10], 3)) +
                                     '\nЭффективная глубина - ' + str(round(all_stat['useful_depth'][10], 2)) +
                                     ' м\nМакс/Мин - ' + str(round(all_stat['max_min'][10], 3)))
    ui.checkBox_sum2.setStyleSheet(k_damp_detect('new', 10))
    ui.checkBox_sum2.setEnabled(True)


def sum3():
    all_stat['name'][11] = text_sum()
    sum_signals(11)
    ui.checkBox_sum3.setText('11 '+text_sum())
    ui.checkBox_sum3.setToolTip('Коэффициент затухания - ' + str(round(all_stat['k_damp'][11], 3)) +
                                     '\nЭффективная глубина - ' + str(round(all_stat['useful_depth'][11], 2)) +
                                     ' м\nМакс/Мин - ' + str(round(all_stat['max_min'][11], 3)))
    ui.checkBox_sum3.setStyleSheet(k_damp_detect('new', 11))
    ui.checkBox_sum3.setEnabled(True)


def sum1_old():
    all_stat_old['name'][13] = text_sum_old()
    sum_signals_old(13)
    ui.checkBox_sum1_old.setText('13 '+text_sum_old())
    ui.checkBox_sum1_old.setToolTip('Коэффициент затухания - ' + str(round(all_stat_old['k_damp'][13], 3)) +
                                  '\nЭффективная глубина - ' + str(round(all_stat_old['useful_depth'][13], 2)) +
                                  ' м\nМакс/Мин - ' + str(round(all_stat_old['max_min'][13], 3)))
    ui.checkBox_sum1_old.setStyleSheet(k_damp_detect('old', 13))
    ui.checkBox_sum1_old.setEnabled(True)


def sum2_old():
    all_stat_old['name'][14] = text_sum_old()
    sum_signals_old(14)
    ui.checkBox_sum2_old.setText('14 '+text_sum_old())
    ui.checkBox_sum2_old.setToolTip('Коэффициент затухания - ' + str(round(all_stat_old['k_damp'][14], 3)) +
                                  '\nЭффективная глубина - ' + str(round(all_stat_old['useful_depth'][14], 2)) +
                                  ' м\nМакс/Мин - ' + str(round(all_stat_old['max_min'][14], 3)))
    ui.checkBox_sum2_old.setStyleSheet(k_damp_detect('old', 14))
    ui.checkBox_sum2_old.setEnabled(True)


def sum3_old():
    all_stat_old['name'][15] = text_sum_old()
    sum_signals_old(15)
    ui.checkBox_sum3_old.setText('15 '+text_sum_old())
    ui.checkBox_sum3_old.setToolTip('Коэффициент затухания - ' + str(round(all_stat_old['k_damp'][15], 3)) +
                                  '\nЭффективная глубина - ' + str(round(all_stat_old['useful_depth'][15], 2)) +
                                  ' м\nМакс/Мин - ' + str(round(all_stat_old['max_min'][15], 3)))
    ui.checkBox_sum3_old.setStyleSheet(k_damp_detect('old', 15))
    ui.checkBox_sum3_old.setEnabled(True)


def text_sum():
    text = 'New: '
    if ui.checkBox_signal1.checkState() == 2:
        text = text + '1 '
    if ui.checkBox_signal2.checkState() == 2:
        text = text + '2 '
    if ui.checkBox_signal3.checkState() == 2:
        text = text + '3 '
    if ui.checkBox_signal4.checkState() == 2:
        text = text + '4 '
    if ui.checkBox_signal5.checkState() == 2:
        text = text + '5 '
    if ui.checkBox_signal6.checkState() == 2:
        text = text + '6 '
    if ui.checkBox_signal7.checkState() == 2:
        text = text + '7 '
    if ui.checkBox_signal8.checkState() == 2:
        text = text + '8 '
    return text


def text_sum_old():
    text = 'Old: '
    if ui.checkBox_signal_old_1.checkState() == 2:
        text = text + '1 '
    if ui.checkBox_signal_old_2.checkState() == 2:
        text = text + '2 '
    if ui.checkBox_signal_old_3.checkState() == 2:
        text = text + '3 '
    if ui.checkBox_signal_old_4.checkState() == 2:
        text = text + '4 '
    if ui.checkBox_signal_old_5.checkState() == 2:
        text = text + '5 '
    if ui.checkBox_signal_old_6.checkState() == 2:
        text = text + '6 '
    if ui.checkBox_signal_old_7.checkState() == 2:
        text = text + '7 '
    if ui.checkBox_signal_old_8.checkState() == 2:
        text = text + '8 '
    if ui.checkBox_signal_old_9.checkState() == 2:
        text = text + '9 '
    if ui.checkBox_signal_old_10.checkState() == 2:
        text = text + '10 '
    if ui.checkBox_signal_old_11.checkState() == 2:
        text = text + '11 '
    if ui.checkBox_signal_old_12.checkState() == 2:
        text = text + '12 '
    return text


def k_damp_detect(device, n_izm):
    if device == 'new':
        if all_stat['k_damp'][n_izm] < 0.3233:
            return 'background: red'
        elif 0.3233 <= all_stat['k_damp'][n_izm] < 0.6806:
            return 'background: yellow'
        elif all_stat['k_damp'][n_izm] >= 0.6806:
            return 'background: green'
    elif device == 'old':
        if all_stat_old['k_damp'][n_izm] < 0.3233:
            return 'background: red'
        elif 0.3233 <= all_stat_old['k_damp'][n_izm] < 0.6806:
            return 'background: yellow'
        elif all_stat_old['k_damp'][n_izm] >= 0.6806:
            return 'background: green'
    # if device == 'new':
    #     if all_stat['k_damp'][n_izm] < 0.65:
    #         return 'background: red'
    #     elif 0.65 <= all_stat['k_damp'][n_izm] < 0.82:
    #         return 'background: yellow'
    #     elif all_stat['k_damp'][n_izm] >= 0.82:
    #         return 'background: green'
    # elif device == 'old':
    #     if all_stat_old['k_damp'][n_izm] < 0.65:
    #         return 'background: red'
    #     elif 0.65 <= all_stat_old['k_damp'][n_izm] < 0.82:
    #         return 'background: yellow'
    #     elif all_stat_old['k_damp'][n_izm] >= 0.82:
    #         return 'background: green'


def calc_int_cement():
    global cement_sig, int_undefine_cement, F, k_damp
    ui.name_cement.setText('измерение цементограммы')
    checkboxes = ui.new_device.findChildren(QCheckBox)
    for item in checkboxes:
        if item.checkState() == 2:
            n_izm = int(item.text()[0:2])
            int_max = calc_max_int(all_signals, all_stat, n_izm)
            if int_min < 86:
                cement_sig = pd.DataFrame(all_signals['depth'].iloc[86:int_max])
            else:
                cement_sig = pd.DataFrame(all_signals['depth'].iloc[int_min:int_max])
            ui.name_cement.setText(item.text())
            k_damp = all_stat['k_damp'][n_izm]
            if n_izm in [1, 2, 5, 6]:
                F = '5'
            elif n_izm in [3, 4, 7, 8]:
                F = '2'
            elif n_izm in [9, 10, 11]:
                F = ''
            if int_min < 86:
                cement_sig['first_sig'] = all_signals[str(n_izm) + '_diff_result'].iloc[86:int_max]
                cement_sig['func'] = all_signals[str(n_izm) + '_func'].iloc[86:int_max]
            else:
                cement_sig['first_sig'] = all_signals[str(n_izm) + '_diff_result'].iloc[int_min:int_max]
                cement_sig['func'] = all_signals[str(n_izm) + '_func'].iloc[int_min:int_max]

    if ui.name_cement.text() == 'измерение цементограммы':
        checkboxes = ui.old_device.findChildren(QCheckBox)
        for item in checkboxes:
            if item.checkState() == 2:
                if int_min_old < 86:
                    cement_sig = pd.DataFrame(all_signals_old['depth'].iloc[86:int_max_old])
                else:
                    cement_sig = pd.DataFrame(all_signals_old['depth'].iloc[int_min_old:int_max_old])
                ui.name_cement.setText(item.text())
                n_izm = int(item.text()[0:2])
                k_damp = all_stat_old['k_damp'][n_izm]
                if n_izm in [1, 2, 5, 6, 9, 10]:
                    F = '5'
                elif n_izm in [3, 4, 7, 8, 11, 12]:
                    F = '2'
                elif n_izm in [13, 14, 15]:
                    F = ''
                if int_min_old < 86:
                    cement_sig['first_sig'] = all_signals_old[str(n_izm) + '_diff_result'].iloc[86:int_max_old]
                    cement_sig['func'] = all_signals_old[str(n_izm) + '_func'].iloc[86:int_max_old]
                else:
                    cement_sig['first_sig'] = all_signals_old[str(n_izm) + '_diff_result'].iloc[int_min_old:int_max_old]
                    cement_sig['func'] = all_signals_old[str(n_izm) + '_func'].iloc[int_min_old:int_max_old]
    cement_sig = cement_sig.reset_index(drop=True)
    cement_sig['corr_sig'] = cement_sig['first_sig'].copy()
    cement_sig['corr_coeff'] = 1

    ui.doubleSpinBox_undefine_max.setMaximum(cement_sig['depth'].max())
    ui.label_int_undefine.setText('Интервалы неопр. цемента:')
    int_undefine_cement = []
    min_value = cement_sig['first_sig'].min()
    max_value = cement_sig['first_sig'].max()
    ui.doubleSpinBox_min_cement.setValue(min_value - (10*(max_value - min_value)/100))
    ui.doubleSpinBox_max_cement.setValue(max_value + (10*(max_value - min_value)/100))
    ui.doubleSpinBox_defect1.setValue(min_value + (max_value - min_value)/3)
    ui.doubleSpinBox_defect2.setValue(min_value + (2*(max_value - min_value)/3))
    ui.doubleSpinBox_x1_line.setValue(cement_sig['depth'].min())
    ui.doubleSpinBox_x2_line.setValue(cement_sig['depth'].min()+10)
    ui.doubleSpinBox_y_line.setValue(ui.doubleSpinBox_defect2.value())
    draw_cement()


def draw_cement():
    """ отрисовка цементограммы в pyqtgraph """
    ui.graphicsView.showGrid(x=False, y=False)  # грид-сетка
    defect1 = ui.doubleSpinBox_defect1.value()
    defect2 = ui.doubleSpinBox_defect2.value()
    ui.graphicsView.clear()

    ui.graphicsView.plot(x=cement_sig['depth'], y=cement_sig['corr_sig'],
                           pen=pg.mkPen(color='#B4FEAE', width=2, dash=[2, 2]))
    d1 = pg.InfiniteLine(pos=defect1, angle=0, pen=pg.mkPen(color='#E59D38', width=0.9, dash=[5, 5]))
    d2 = pg.InfiniteLine(pos=defect2, angle=0, pen=pg.mkPen(color='#FF1705', width=0.9, dash=[5, 5]))
    ui.graphicsView.addItem(d1)
    ui.graphicsView.addItem(d2)

    if ui.checkBox_corr_mode.checkState() == 2:
        ui.graphicsView.plot(x=cement_sig['depth'], y=cement_sig['first_sig'],
                               pen=pg.mkPen(color='#FF3D70', width=2, dash=[8, 3, 2, 3]))
        line_x1 = ui.doubleSpinBox_x1_line.value()
        line_x2 = ui.doubleSpinBox_x2_line.value()
        line_y = ui.doubleSpinBox_y_line.value()
        lx1 = pg.InfiniteLine(pos=line_x1, angle=90, pen=pg.mkPen(width=0.7, dash=[8, 2]))
        lx2 = pg.InfiniteLine(pos=line_x2, angle=90, pen=pg.mkPen(width=0.7, dash=[8, 2]))
        ly = pg.InfiniteLine(pos=line_y, angle=0, pen=pg.mkPen(width=0.7, dash=[8, 2]))
        ui.graphicsView.addItem(lx1)
        ui.graphicsView.addItem(lx2)
        ui.graphicsView.addItem(ly)


def corr_sig():
    x1_index = cement_sig.index[
        cement_sig['depth'] == get_nearest_value(cement_sig['depth'], ui.doubleSpinBox_x1_line.value())].tolist()[0]
    x2_index = cement_sig.index[
        cement_sig['depth'] == get_nearest_value(cement_sig['depth'], ui.doubleSpinBox_x2_line.value())].tolist()[0]
    cement_sig['plus_sig'] = cement_sig['first_sig'] + abs(cement_sig['first_sig'].min())
    corr_coeff_max = (ui.doubleSpinBox_y_line.value() + abs(cement_sig['first_sig'].min()))/\
                 cement_sig['plus_sig'].iloc[x1_index:x2_index+1].max()
    xmax_index = cement_sig.index[cement_sig['plus_sig'] == get_nearest_value(cement_sig['plus_sig'],
                                                                              cement_sig['plus_sig'].iloc[
                                                                              x1_index:x2_index + 1].max())].tolist()[0]
    if ui.checkBox_lin_cor.checkState() == 2:
        interp_method = 'linear'
    else:
        interp_method = 'quadratic'

    if x1_index == xmax_index:
        corr_coeff = interp1d([x1_index, x2_index], [corr_coeff_max, 1], kind='linear')
        print(x1_index, xmax_index, x2_index)
    elif xmax_index == x2_index:
        corr_coeff = interp1d([x1_index, x2_index], [1, corr_coeff_max], kind='linear')
        print(x1_index, xmax_index, x2_index)
    else:
        corr_coeff = interp1d([x1_index, xmax_index, x2_index], [1, corr_coeff_max, 1], kind=interp_method)
        print(x1_index, xmax_index, x2_index)

    cement_sig['corr_coeff'].iloc[x1_index:x2_index+1] = corr_coeff(range(x1_index, x2_index+1))
    cement_sig['plus_sig'] = cement_sig['plus_sig']*cement_sig['corr_coeff']
    cement_sig['corr_sig'] = cement_sig['plus_sig'] - abs(cement_sig['first_sig'].min())
    draw_cement()


def corr_sig_bottom():
    x1_index = cement_sig.index[
        cement_sig['depth'] == get_nearest_value(cement_sig['depth'], ui.doubleSpinBox_x1_line.value())].tolist()[0]
    x2_index = cement_sig.index[
        cement_sig['depth'] == get_nearest_value(cement_sig['depth'], ui.doubleSpinBox_x2_line.value())].tolist()[0]
    cement_sig['plus_sig'] = cement_sig['first_sig'] + abs(cement_sig['first_sig'].min())
    corr_coeff_max = (ui.doubleSpinBox_y_line.value() + abs(cement_sig['first_sig'].min())) / \
                     cement_sig['plus_sig'].iloc[x1_index:x2_index+1].min()
    xmax_index = cement_sig.index[cement_sig['plus_sig'] == get_nearest_value(cement_sig['plus_sig'],
                                                                              cement_sig['plus_sig'].iloc[
                                                                              x1_index:x2_index+1].min())].tolist()[0]
    if ui.checkBox_lin_cor.checkState() == 2:
        interp_method = 'linear'
    else:
        interp_method = 'quadratic'
    if x1_index == xmax_index:
        corr_coeff = interp1d([x1_index, x2_index], [corr_coeff_max, 1], kind='linear')
        print(x1_index, xmax_index, x2_index)
    elif xmax_index == x2_index:
        corr_coeff = interp1d([x1_index, x2_index], [1, corr_coeff_max], kind='linear')
        print(x1_index, xmax_index, x2_index)
    else:
        corr_coeff = interp1d([x1_index, xmax_index, x2_index], [1, corr_coeff_max, 1], kind=interp_method)
        print(x1_index, xmax_index, x2_index)

    cement_sig['corr_coeff'].iloc[x1_index:x2_index+1] = corr_coeff(range(x1_index, x2_index+1))
    cement_sig['plus_sig'] = cement_sig['plus_sig'] * cement_sig['corr_coeff']
    cement_sig['corr_sig'] = cement_sig['plus_sig'] - abs(cement_sig['first_sig'].min())
    draw_cement()


def fix_corr():
    cement_sig['first_sig'] = cement_sig['corr_sig']
    draw_cement()


def plus_func():
    cement_sig['corr_sig'] = cement_sig['first_sig'] + cement_sig['func']
    draw_cement()


def add_undefine():
    min_undefine = ui.doubleSpinBox_undefine_min.value()
    max_undefine = ui.doubleSpinBox_undefine_max.value()
    minund_i = cement_sig.index[cement_sig['depth'] == get_nearest_value(cement_sig['depth'], min_undefine)].tolist()[0]
    maxund_i = cement_sig.index[cement_sig['depth'] == get_nearest_value(cement_sig['depth'], max_undefine)].tolist()[0]
    int_undefine_cement.append([minund_i, maxund_i])
    text_int = ui.label_int_undefine.text()
    text_int = text_int + '\n' + str(min_undefine) + ' - ' + str(max_undefine) + ' м.'
    ui.label_int_undefine.setText(text_int)


def cementogramma():
    global options
    min_cement = ui.doubleSpinBox_min_cement.value()
    max_cement = ui.doubleSpinBox_max_cement.value()
    defect1 = ui.doubleSpinBox_defect1.value()
    defect2 = ui.doubleSpinBox_defect2.value()
    cement_sig['quality'] = 0
    begin = []
    end = []
    try:
        if cement_sig['corr_sig'][0] > defect1:
            begin.append(0)
        for n, i in enumerate(cement_sig['corr_sig']):
            if 0 < n < (len(cement_sig['corr_sig']) - 1):
                if i >= defect1:
                    if cement_sig['corr_sig'][n-1] < defect1:
                        begin.append(n)
                    if cement_sig['corr_sig'][n+1] < defect1:
                        end.append(n+1)
        if cement_sig['corr_sig'][len(cement_sig['corr_sig'])-1] >= defect1:
            end.append(len(cement_sig['corr_sig']))
        for i in range(len(begin)):
            if cement_sig['corr_sig'].iloc[begin[i]:end[i]].max() >= defect2:
                cement_sig['quality'].iloc[begin[i]:end[i]] = 2
            else:
                cement_sig['quality'].iloc[begin[i]:end[i]] = 1
        for i in int_undefine_cement:
            cement_sig['quality'].iloc[i[0]:i[1]] = 3
        ui.label_info.setText('Готово!')
        ui.label_info.setStyleSheet('color:green')
    except IndexError:
        ui.label_info.setText('Ошибка расчета интервалов дефектов. Поробуйте уменьшить максимальную глубину. '
                                'Например на 1 метр.')
        ui.label_info.setStyleSheet('color:red')

    def draw_cement_int(ax, up):
        ax.axvline(x=defect1, linewidth=0.5, color='black', linestyle=':')
        ax.axvline(x=defect2, linewidth=0.5, color='black', linestyle=':')
        plt.fill_betweenx(cement_sig['depth'], max_cement, cement_sig['corr_sig'], where=cement_sig['quality'] >= 1,
                          hatch='//', facecolor='#EDEDED')
        plt.fill_betweenx(cement_sig['depth'], max_cement, cement_sig['corr_sig'], where=cement_sig['quality'] >= 2,
                          hatch='////', facecolor='#BDBDBD')
        plt.fill_betweenx(cement_sig['depth'], max_cement, cement_sig['corr_sig'], where=cement_sig['quality'] >= 3,
                          hatch='...', facecolor='white', edgecolor='grey')
        plt.fill_betweenx(cement_sig['depth'], min_cement, cement_sig['corr_sig'], where=cement_sig['quality'] >= 1,
                          hatch='//', facecolor='#EDEDED')
        plt.fill_betweenx(cement_sig['depth'], min_cement, cement_sig['corr_sig'], where=cement_sig['quality'] >= 2,
                          hatch='////', facecolor='#BDBDBD')
        plt.fill_betweenx(cement_sig['depth'], min_cement, cement_sig['corr_sig'], where=cement_sig['quality'] >= 3,
                          hatch='...', facecolor='white', edgecolor='grey')
        plt.plot(cement_sig['corr_sig'], cement_sig['depth'], 'black')
        plt.ylim(up, up + 100)
        plt.xlim(min_cement, max_cement)
        plt.title('ЧЦ', fontsize=9)
        plt.title('КЦ', loc='left', fontsize=9)
        plt.title('ПЦ', loc='right', fontsize=9)
        ax.get_xaxis().set_visible(False)
        plt.yticks(np.arange(up, up + 101, 5))
        plt.grid(axis='y', color='black', linestyle=':', linewidth=0.5)
        ax.invert_yaxis()

    fig = plt.figure(figsize=(6.3, 8.45))
    ax1 = plt.subplot(141)
    draw_cement_int(ax1, 0)
    ax2 = plt.subplot(142)
    draw_cement_int(ax2, 100)
    ax3 = plt.subplot(143)
    draw_cement_int(ax3, 200)
    ax4 = plt.subplot(144)
    draw_cement_int(ax4, 300)

    fig.tight_layout()
    fig.savefig('temp.png', dpi=200)
    fig.suptitle(ui.name_cement.text())
    fig.tight_layout()
    fig.show()

    if cement_sig['depth'].max() > 400:
        fig = plt.figure(figsize=(6.3, 8.45))
        ax1 = plt.subplot(141)
        draw_cement_int(ax1, 400)
        ax2 = plt.subplot(142)
        draw_cement_int(ax2, 500)
        ax3 = plt.subplot(143)
        draw_cement_int(ax3, 600)
        ax4 = plt.subplot(144)
        draw_cement_int(ax4, 700)

        fig.tight_layout()
        fig.savefig('temp1.png', dpi=200)
        fig.suptitle(ui.name_cement.text())
        fig.tight_layout()
        fig.show()

    if cement_sig['depth'].max() > 800:
        fig = plt.figure(figsize=(6.3, 8.45))
        ax1 = plt.subplot(141)
        draw_cement_int(ax1, 800)
        ax2 = plt.subplot(142)
        draw_cement_int(ax2, 900)
        ax3 = plt.subplot(143)
        draw_cement_int(ax3, 1000)
        ax4 = plt.subplot(144)
        draw_cement_int(ax4, 1100)

        fig.tight_layout()
        fig.savefig('temp2.png', dpi=200)
        fig.suptitle(ui.name_cement.text())
        fig.tight_layout()
        fig.show()

    cement_sig.to_csv('result_cement.txt', sep='\t')

    # словарь настроек для сохранения
    options = {}
    options['doubleSpinBox_int_min'] = ui.doubleSpinBox_int_min.value()
    options['doubleSpinBox_int_max'] = ui.doubleSpinBox_int_max.value()
    options['doubleSpinBox_min_cement'] = min_cement
    options['doubleSpinBox_max_cement'] = max_cement
    options['doubleSpinBox_defect1'] = defect1
    options['doubleSpinBox_defect2'] = defect2
    options['name_signal'] = ui.name_cement.text()
    options['signal'] = cement_sig['corr_sig'].tolist()


def check_int_corr():
    """ проверка интервала коррекции"""
    ui.doubleSpinBox_x1_line.setMaximum(ui.doubleSpinBox_x2_line.value())
    ui.doubleSpinBox_x2_line.setMinimum(ui.doubleSpinBox_x1_line.value())
    draw_cement()


def check_int_defect():
    """ проверка интервала уровней дефектов"""
    ui.doubleSpinBox_defect1.setMaximum(ui.doubleSpinBox_defect2.value())
    ui.doubleSpinBox_defect2.setMinimum(ui.doubleSpinBox_defect1.value())
    draw_cement()


def check_int_undefine():
    """ проверка интервала неопределенного цемента"""
    ui.doubleSpinBox_undefine_min.setMaximum(ui.doubleSpinBox_undefine_max.value())
    ui.doubleSpinBox_undefine_max.setMinimum(ui.doubleSpinBox_undefine_min.value())


def cement_temp():
    # Работа с шаблоном Word
    doc = docx.Document('template.docx')
    table = doc.tables[0]
    ui.lineEdit_pipe_diameter.setText(table.rows[1].cells[0].paragraphs[0].runs[3].text)
    ui.lineEdit_client.setText(table.rows[2].cells[0].paragraphs[0].runs[1].text)
    ui.lineEdit_client2.setText(table.rows[3].cells[0].paragraphs[0].text)
    ui.lineEdit_interpreter.setText(table.rows[9].cells[2].paragraphs[0].text)


def save_cement():
    """
    Сохранение результирующей цементограммы в отчет в документ Word по шаблону 'template.docx'
    """
    doc = docx.Document('template.docx')
    table = doc.tables[0]
    table2 = doc.tables[1]

    # номер скважины
    table.rows[0].cells[0].paragraphs[0].runs[3].text = ui.lineEdit_number_skv.text()
    doc.paragraphs[4].runs[2].text = ui.lineEdit_number_skv.text()
    doc.paragraphs[9].runs[5].text = ui.lineEdit_number_skv.text()

    # тип трубы
    table.rows[1].cells[0].paragraphs[0].runs[0].text = ui.comboBox_pipe_type.currentText()
    if ui.comboBox_pipe_type.currentText() == 'Кондуктор':
        tube = 'кондуктора'
    elif ui.comboBox_pipe_type.currentText() == 'Колонна':
        tube = 'колонны'
    elif ui.comboBox_pipe_type.currentText() == 'Тех. колонна':
        tube = 'тех. колонны'
    elif ui.comboBox_pipe_type.currentText() == 'Направление':
        tube = 'направления'
    else:
        tube = 'кондуктора'
    doc.paragraphs[9].runs[3].text = tube

    # диаметр трубы
    table.rows[1].cells[0].paragraphs[0].runs[3].text = ui.lineEdit_pipe_diameter.text()
    doc.paragraphs[9].runs[7].text = ui.lineEdit_pipe_diameter.text()

    # заказчик
    table.rows[2].cells[0].paragraphs[0].runs[1].text = ui.lineEdit_client.text()
    doc.paragraphs[6].runs[1].text = ui.lineEdit_client.text()
    table.rows[3].cells[0].paragraphs[0].clear()
    doc.paragraphs[7].clear()
    client2 = table.rows[3].cells[0].paragraphs[0].add_run(ui.lineEdit_client2.text())
    client2.font.size = Pt(11)
    client2.font.name = 'Arial'
    client2 = doc.paragraphs[7].add_run(ui.lineEdit_client2.text())
    client2.font.size = Pt(12)
    client2.font.name = 'Times New Roman'

    # интервал исследования
    table.rows[1].cells[1].paragraphs[0].runs[1].text = str(round(cement_sig['depth'].min(), 1))
    table.rows[1].cells[1].paragraphs[0].runs[5].text = str(float(ui.doubleSpinBox_int_max.value()))
    doc.paragraphs[9].runs[10].text = str(round(cement_sig['depth'].min(), 1))
    doc.paragraphs[9].runs[12].text = str(round(cement_sig['depth'].max(), 1))
    int_study = cement_sig['depth'].max() - cement_sig['depth'].min()

    # дата исследования
    table.rows[2].cells[1].paragraphs[0].runs[1].text = ui.dateEdit.text()

    # интерпретатор
    table.rows[9].cells[2].paragraphs[0].clear()
    interpreter = table.rows[9].cells[2].paragraphs[0].add_run(ui.lineEdit_interpreter.text())
    interpreter.font.size = Pt(11)
    interpreter.font.name = 'Times New Roman'
    table2.rows[1].cells[1].paragraphs[0].clear()
    interpreter = table2.rows[1].cells[1].paragraphs[0].add_run(ui.lineEdit_interpreter.text())
    interpreter.font.size = Pt(11)
    interpreter.font.name = 'Times New Roman'

    # цементограмма
    table.rows[6].cells[0]._element.clear_content()
    table.rows[6].cells[0].add_paragraph().add_run().add_picture('temp.png')
    table.rows[6].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # частота
    table.rows[7].cells[2].paragraphs[0].runs[5].text = F

    # Интервалы дефектов
    n_row = 1
    h_def = 0
    if cement_sig['quality'][0] > 0:
        table2.rows[0].cells[0].tables[0].add_row()
        note = table2.rows[0].cells[0].tables[0].rows[n_row].cells[0].paragraphs[0].add_run(str(round(cement_sig['depth'][0], 1)))
        note.font.size = Pt(11)
        note.font.name = 'Times New Roman'
        table2.rows[0].cells[0].tables[0].rows[n_row].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, quality in enumerate(cement_sig['quality']):
        if 0 < i < (len(cement_sig['quality']) - 1):
            if quality > 0:
                if cement_sig['quality'][i-1] != quality:
                    table2.rows[0].cells[0].tables[0].add_row()
                    note = table2.rows[0].cells[0].tables[0].rows[n_row].cells[0].paragraphs[0].add_run(
                        str(round(cement_sig['depth'][i], 1)))
                    note.font.size = Pt(11)
                    note.font.name = 'Times New Roman'
                    table2.rows[0].cells[0].tables[0].rows[n_row].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cement_sig['quality'][i + 1] != quality:
                    note = table2.rows[0].cells[0].tables[0].rows[n_row].cells[1].paragraphs[0].add_run(
                        str(round(cement_sig['depth'][i], 1)))
                    note.font.size = Pt(11)
                    note.font.name = 'Times New Roman'
                    table2.rows[0].cells[0].tables[0].rows[n_row].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    h = float(table2.rows[0].cells[0].tables[0].rows[n_row].cells[1].paragraphs[0].text) - \
                        float(table2.rows[0].cells[0].tables[0].rows[n_row].cells[0].paragraphs[0].text)
                    h_def = h_def + h
                    note = table2.rows[0].cells[0].tables[0].rows[n_row].cells[2].paragraphs[0].add_run(str(round(h, 1)))
                    note.font.size = Pt(11)
                    note.font.name = 'Times New Roman'
                    table2.rows[0].cells[0].tables[0].rows[n_row].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if quality == 1:
                        defect = 'частичный цемент'
                    elif quality == 2:
                        defect = 'плохой цемент'
                    elif quality == 3:
                        defect = 'неопределённый цемент'
                    note = table2.rows[0].cells[0].tables[0].rows[n_row].cells[3].paragraphs[0].add_run(defect)
                    note.font.size = Pt(11)
                    note.font.name = 'Times New Roman'
                    table2.rows[0].cells[0].tables[0].rows[n_row].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    n_row += 1
    if cement_sig['quality'][len(cement_sig['quality']) - 1] > 0:
        note = table2.rows[0].cells[0].tables[0].rows[n_row].cells[1].paragraphs[0].add_run(
            str(round(cement_sig['depth'][len(cement_sig['quality']) - 1], 1)))
        note.font.size = Pt(11)
        note.font.name = 'Times New Roman'
        table2.rows[0].cells[0].tables[0].rows[n_row].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        h = float(table2.rows[0].cells[0].tables[0].rows[n_row].cells[1].paragraphs[0].text) - \
            float(table2.rows[0].cells[0].tables[0].rows[n_row].cells[0].paragraphs[0].text)
        h_def = h_def + h
        note = table2.rows[0].cells[0].tables[0].rows[n_row].cells[2].paragraphs[0].add_run(str(round(h, 1)))
        note.font.size = Pt(11)
        note.font.name = 'Times New Roman'
        table2.rows[0].cells[0].tables[0].rows[n_row].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cement_sig['quality'][len(cement_sig['quality']) - 1] == 1:
            defect = 'частичный цемент'
        elif cement_sig['quality'][len(cement_sig['quality']) - 1] == 2:
            defect = 'плохой цемент'
        elif cement_sig['quality'][len(cement_sig['quality']) - 1] == 3:
            defect = 'неопределённый цемент'
        note = table2.rows[0].cells[0].tables[0].rows[n_row].cells[3].paragraphs[0].add_run(defect)
        note.font.size = Pt(11)
        note.font.name = 'Times New Roman'
        table2.rows[0].cells[0].tables[0].rows[n_row].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # коэффициент затухания
    k_def = 1 - h_def / int_study
    table.rows[7].cells[0].paragraphs[0].runs[7].text = str(round(k_def, 1))

    if cement_sig['depth'].max() > 400:
        # копирование таблицы
        new_tbl = deepcopy(table._tbl)
        table._tbl.addnext(new_tbl)
        # цементограмма
        table3 = doc.tables[1]
        table3.rows[6].cells[0]._element.clear_content()
        table3.rows[6].cells[0].add_paragraph().add_run().add_picture('temp1.png')
        table3.rows[6].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table3.rows[0].cells[3].paragraphs[0].runs[0].text = 'лист 2'

    if cement_sig['depth'].max() > 800:
        new_tbl2 = deepcopy(table3._tbl)
        table3._tbl.addnext(new_tbl2)
        # цементограмма
        table4 = doc.tables[2]
        table4.rows[6].cells[0]._element.clear_content()
        table4.rows[6].cells[0].add_paragraph().add_run().add_picture('temp2.png')
        table4.rows[6].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table4.rows[0].cells[3].paragraphs[0].runs[0].text = 'лист 3'
    try:
        doc.save('цементограмма_по_скв_' + ui.lineEdit_number_skv.text() + '.docx')
        ui.label_info.setText('Результат сохранён в файл "цементограмма_по_скв_' + ui.lineEdit_number_skv.text() +
                                '.docx"')
        ui.label_info.setStyleSheet('color:green')
    except PermissionError:
        ui.label_info.setText('Сохранение невозможно! Закройте файл "цементограмма_по_скв_' +
                                ui.lineEdit_number_skv.text() + '.docx"')
        ui.label_info.setStyleSheet('color:red')

    with open('options_' + ui.lineEdit_number_skv.text() + '.vac', "wb") as f:
        pickle.dump(options, f)


def draw_total_cement(device, n_izm, n_graph, fig, count_sig):
    if device is all_signals_old:
        int_min_self = int_min_old
        int_max_self = int_max_old
    else:
        int_min_self = int_min
        int_max_self = int_max
    signal = pd.DataFrame(device['depth'].iloc[int_min_self:int_max_self].copy())
    signal['curve'] = device[str(n_izm) + '_diff_result'].iloc[int_min_self:int_max_self].copy()
    signal = signal.reset_index(drop=True)
    max_depth = signal['depth'].max()
    min_depth = signal['depth'].min()
    min_value = signal['curve'].min()
    max_value = signal['curve'].max()
    min_cement = min_value - (10 * (max_value - min_value) / 100)
    max_cement = max_value + (10 * (max_value - min_value) / 100)
    defect1 = min_value + (max_value - min_value) / 3
    defect2 = min_value + (2 * (max_value - min_value) / 3)
    # defect1 = 1
    # defect2 = 5
    signal['quality'] = 0
    begin = []
    end = []
    try:
        if signal['curve'][0] > defect1:
            begin.append(0)
        for n, i in enumerate(signal['curve']):
            if 0 < n < (len(signal['curve']) - 1):
                if i >= defect1:
                    if signal['curve'][n - 1] < defect1:
                        begin.append(n)
                    if signal['curve'][n + 1] < defect1:
                        end.append(n + 1)
        if signal['curve'][len(signal['curve']) - 1] >= defect1:
            end.append(len(signal['curve']))
        for i in range(len(begin)):
            if signal['curve'].iloc[begin[i]:end[i]].max() >= defect2:
                signal['quality'].iloc[begin[i]:end[i]] = 2
            else:
                signal['quality'].iloc[begin[i]:end[i]] = 1
    except IndexError:
        ui.label_info.setText('Ошибка расчета интервалов дефектов. Поробуйте уменьшить максимальную глубину. '
                                'Например на 1 метр.')
        ui.label_info.setStyleSheet('color:red')
    ax = fig.add_subplot(1, count_sig, n_graph)
    ax.axvline(x=defect1, linewidth=0.5, color='black', linestyle=':')
    ax.axvline(x=defect2, linewidth=0.5, color='black', linestyle=':')
    ax.fill_betweenx(signal['depth'], max_cement, signal['curve'], where=signal['quality'] >= 1, hatch='//',
                     facecolor='#EDEDED')
    ax.fill_betweenx(signal['depth'], max_cement, signal['curve'], where=signal['quality'] >= 2, hatch='////',
                     facecolor='#BDBDBD')
    ax.fill_betweenx(signal['depth'], min_cement, signal['curve'], where=signal['quality'] >= 1, hatch='//',
                     facecolor='#EDEDED')
    ax.fill_betweenx(signal['depth'], min_cement, signal['curve'], where=signal['quality'] >= 2, hatch='////',
                     facecolor='#BDBDBD')
    ax.plot(signal['curve'], signal['depth'], 'black')
    plt.ylim(min_depth, max_depth)
    plt.xlim(min_cement, max_cement)
    if device is all_signals:
        plt.title(all_stat['name'][n_izm][-35:-23])
    if device is all_signals_old:
        plt.title(all_stat_old['name'][n_izm])
    ax.get_xaxis().set_visible(False)
    plt.yticks(np.arange(min_depth, max_depth, 5))
    ax.grid(axis='y', color='black', linestyle=':', linewidth=0.5)
    ax.invert_yaxis()
    ui.label_info.setText('Готово!')
    ui.label_info.setStyleSheet('color:green')


def all_cement():
    fig = plt.figure(figsize=(12, 9))
    for i in range(1, 9):
        draw_total_cement(all_signals, i, i, fig, 8)
    fig.tight_layout()
    fig.show()


def all_cement_old():
    fig = plt.figure(figsize=(18, 9))
    for i in range(1, 13):
        draw_total_cement(all_signals_old, i, i, fig, 12)
    fig.tight_layout()
    fig.show()


def choice_cement():
    choice_signal_new = []
    choice_signal_old = []
    checkboxes = ui.new_device.findChildren(QCheckBox)
    for item in checkboxes:
        if item.checkState() == 2:
            choice_signal_new.append(int(item.text()[0:2]))
    checkboxes = ui.old_device.findChildren(QCheckBox)
    for item in checkboxes:
        if item.checkState() == 2:
            choice_signal_old.append(int(item.text()[0:2]))
    count_choice_signal = len(choice_signal_new) + len(choice_signal_old)
    fig = plt.figure(figsize=(1.5*count_choice_signal, 9))
    for n, i in enumerate(choice_signal_new):
        draw_total_cement(all_signals, i, n+1, fig, count_choice_signal)
    for n, i in enumerate(choice_signal_old):
        draw_total_cement(all_signals_old, i, len(choice_signal_new)+n+1,  fig, count_choice_signal)
    fig.tight_layout()
    fig.show()


def background_white():
    """переключение цвета фона графика"""
    if ui.checkBox_background_white.checkState() == 2:
        ui.graphicsView.setBackground('w')
    else:
        ui.graphicsView.setBackground('k')


def open_options():
    """Загрузка настроек обработки"""
    file_name = QFileDialog.getOpenFileName(filter='*.vac')
    with open(file_name[0], 'rb') as f:
        load_options = pickle.load(f)

    ui.doubleSpinBox_int_min.setValue(load_options['doubleSpinBox_int_min'])
    ui.doubleSpinBox_int_max.setValue(load_options['doubleSpinBox_int_max'])

    if load_options['name_signal'][3:8] == 'Old: ':
        sum_sig = load_options['name_signal'][8:-1].split(' ')
        for i, item in enumerate(sum_sig):
            sum_sig[i] = int(item)
        checkboxes = ui.old_device.findChildren(QCheckBox)
        for item in checkboxes:
            if int(item.text()[0:2]) in sum_sig:
                item.setCheckState(2)
        sum1_old()
        ui.checkBox_sum1_old.setCheckState(2)
        for item in checkboxes:
            if int(item.text()[0:2]) in sum_sig:
                item.setCheckState(0)
    elif load_options['name_signal'][3:8] == 'New: ':
        sum_sig = load_options['name_signal'][8:-1].split(' ')
        for i, item in enumerate(sum_sig):
            sum_sig[i] = int(item)
        checkboxes = ui.new_device.findChildren(QCheckBox)
        for item in checkboxes:
            if int(item.text()[0:2]) in sum_sig:
                item.setCheckState(2)
        sum1()
        ui.checkBox_sum1.setCheckState(2)
        for item in checkboxes:
            if int(item.text()[0:2]) in sum_sig:
                item.setCheckState(0)
    else:
        checkboxes = ui.new_device.findChildren(QCheckBox)
        for item in checkboxes:
            item.setCheckState(0)
            if item.text() == load_options['name_signal']:
                item.setCheckState(2)
        checkboxes = ui.old_device.findChildren(QCheckBox)
        for item in checkboxes:
            item.setCheckState(0)
            if item.text() == load_options['name_signal']:
                item.setCheckState(2)
    calc()
    calc_int_cement()
    cement_sig['corr_sig'] = load_options['signal']
    ui.doubleSpinBox_min_cement.setValue(load_options['doubleSpinBox_min_cement'])
    ui.doubleSpinBox_max_cement.setValue(load_options['doubleSpinBox_max_cement'])
    ui.doubleSpinBox_defect1.setValue(load_options['doubleSpinBox_defect1'])
    ui.doubleSpinBox_defect2.setValue(load_options['doubleSpinBox_defect2'])
    ui.checkBox_corr_mode.setCheckState(2)


def set_int(tab):
    ui.doubleSpinBox_int_max.setMaximum(tab['depth'].max())  # устанавливаем в спинбоксы интервала сигнала
    ui.doubleSpinBox_int_min.setMaximum(tab['depth'].max())  # максимально допустимые значения из файла
    ui.doubleSpinBox_int_max.setValue(tab['depth'].max())  # устанавливаем в спинбоксы значения мин и макс
    ui.doubleSpinBox_int_min.setValue(tab['depth'].min())


def calc_max_int(tab1, tab2, i):
    if tab2['max_h'][i] < ui.doubleSpinBox_int_max.value():
        res = tab1.index[tab1['depth'] == get_nearest_value(tab1['depth'], tab2['max_h'][i])].tolist()[0]
        return res
    else:
        res = tab1.index[tab1['depth'] == get_nearest_value(tab1['depth'],
                                                                ui.doubleSpinBox_int_max.value())].tolist()[0]
        return res


ui.Button_direct.clicked.connect(open_dir)
ui.Button_direct_old.clicked.connect(open_dir_old)
ui.Button_open_sig.clicked.connect(open_sig)
ui.Button_open_sig_old.clicked.connect(open_sig_old)

ui.checkBox_signal1.stateChanged.connect(choice_signal)
ui.checkBox_signal2.stateChanged.connect(choice_signal)
ui.checkBox_signal3.stateChanged.connect(choice_signal)
ui.checkBox_signal4.stateChanged.connect(choice_signal)
ui.checkBox_signal5.stateChanged.connect(choice_signal)
ui.checkBox_signal6.stateChanged.connect(choice_signal)
ui.checkBox_signal7.stateChanged.connect(choice_signal)
ui.checkBox_signal8.stateChanged.connect(choice_signal)
ui.checkBox_noise.stateChanged.connect(choice_signal)
ui.checkBox_origin_sig.stateChanged.connect(choice_signal)
ui.checkBox_envelop.stateChanged.connect(choice_signal)
ui.checkBox_func.stateChanged.connect(choice_signal)
ui.checkBox_mean.stateChanged.connect(choice_signal)
ui.checkBox_norm.stateChanged.connect(choice_signal)
ui.checkBox_rel_ampl.stateChanged.connect(choice_signal)
ui.checkBox_diff_result.stateChanged.connect(choice_signal)

ui.checkBox_signal_old_1.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_2.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_3.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_4.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_5.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_6.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_7.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_8.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_9.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_10.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_11.stateChanged.connect(choice_signal)
ui.checkBox_signal_old_12.stateChanged.connect(choice_signal)
ui.checkBox_noise_old.stateChanged.connect(choice_signal)
ui.checkBox_envelop_old.stateChanged.connect(choice_signal)
ui.checkBox_func_old.stateChanged.connect(choice_signal)
ui.checkBox_mean_old.stateChanged.connect(choice_signal)
ui.checkBox_norm_old.stateChanged.connect(choice_signal)
ui.checkBox_rel_ampl_old.stateChanged.connect(choice_signal)
ui.checkBox_diff_result_old.stateChanged.connect(choice_signal)

ui.spinBox_mean_win.valueChanged.connect(change_mean_win)
ui.doubleSpinBox_coeff_norm.valueChanged.connect(change_mean_win)
ui.spinBox_coeff_dif_res.valueChanged.connect(change_mean_win)

ui.spinBox_mean_win_old.valueChanged.connect(change_mean_win_old)
ui.doubleSpinBox_coeff_norm_old.valueChanged.connect(change_mean_win_old)
ui.spinBox_coeff_dif_res_old.valueChanged.connect(change_mean_win_old)

ui.doubleSpinBox_coeff_func.valueChanged.connect(change_func)
ui.doubleSpinBox_kA.valueChanged.connect(change_func)
ui.doubleSpinBox_kB.valueChanged.connect(change_func)
ui.doubleSpinBox_coeff_func_old.valueChanged.connect(change_func_old)
ui.doubleSpinBox_kA_old.valueChanged.connect(change_func_old)
ui.doubleSpinBox_kB_old.valueChanged.connect(change_func_old)

ui.doubleSpinBox_int_max.valueChanged.connect(check_int_sig)
ui.doubleSpinBox_int_max.valueChanged.connect(check_int_sig)
ui.pushButton_int.clicked.connect(calc)

ui.pushButton_sum1.clicked.connect(sum1)
ui.pushButton_sum2.clicked.connect(sum2)
ui.pushButton_sum3.clicked.connect(sum3)
ui.checkBox_sum1.stateChanged.connect(choice_signal)
ui.checkBox_sum2.stateChanged.connect(choice_signal)
ui.checkBox_sum3.stateChanged.connect(choice_signal)

ui.pushButton_sum1_old.clicked.connect(sum1_old)
ui.pushButton_sum2_old.clicked.connect(sum2_old)
ui.pushButton_sum3_old.clicked.connect(sum3_old)
ui.checkBox_sum1_old.stateChanged.connect(choice_signal)
ui.checkBox_sum2_old.stateChanged.connect(choice_signal)
ui.checkBox_sum3_old.stateChanged.connect(choice_signal)

ui.pushButton_int_cement.clicked.connect(calc_int_cement)
ui.doubleSpinBox_defect1.valueChanged.connect(check_int_defect)
ui.doubleSpinBox_defect2.valueChanged.connect(check_int_defect)
ui.doubleSpinBox_x1_line.valueChanged.connect(check_int_corr)
ui.doubleSpinBox_x2_line.valueChanged.connect(check_int_corr)
ui.doubleSpinBox_y_line.valueChanged.connect(draw_cement)
ui.pushButton_corr_sig.clicked.connect(corr_sig)
ui.pushButton_corr_sig_bottom.clicked.connect(corr_sig_bottom)
ui.pushButton_fix_corr.clicked.connect(fix_corr)
ui.pushButton_plus_func.clicked.connect(plus_func)

ui.doubleSpinBox_undefine_min.valueChanged.connect(check_int_undefine)
ui.doubleSpinBox_undefine_max.valueChanged.connect(check_int_undefine)
ui.pushButton_add_undefine.clicked.connect(add_undefine)

ui.pushButton_cement.clicked.connect(cementogramma)

ui.checkBox_corr_mode.stateChanged.connect(draw_cement)
ui.pushButton_from_temp.clicked.connect(cement_temp)
ui.pushButton_save_cement.clicked.connect(save_cement)
ui.pushButton_open_options.clicked.connect(open_options)

ui.pushButton_all_cement.clicked.connect(all_cement)
ui.pushButton_all_cement_old.clicked.connect(all_cement_old)
ui.pushButton_choice_cement.clicked.connect(choice_cement)

ui.checkBox_background_white.stateChanged.connect(background_white)

sys.exit(app.exec_())
